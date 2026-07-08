from importlib import metadata
import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
# pyrefly: ignore [missing-import]
import chromadb
# pyrefly: ignore [missing-import]
import fitz  # PyMuPDF
import csv
# pyrefly: ignore [missing-import]
import docx
# pyrefly: ignore [missing-import]
import httpx
# pyrefly: ignore [missing-import]
from openai import OpenAI
from core.config import settings, ROOT_DIR
from core.embeddings import get_embedding_function, get_cached_embedding

logger = logging.getLogger(__name__)


def get_openai_client() -> Optional[OpenAI]:
    api_key = settings.OPENROUTER_API_KEY or settings.OPENAI_API_KEY
    if not api_key:
        return None
    return OpenAI(
        base_url=settings.OPENAI_BASE_URL,
        api_key=api_key,
    )

# Constants
CHROMA_PATH = ROOT_DIR / "chroma_db"

# Initialize ChromaDB Client (Persistent)
chroma_client = chromadb.PersistentClient(path=str(CHROMA_PATH))

# Create/Get collection configured with cosine similarity
# We base the collection name on the active model
model_slug = re.sub(r'[^a-zA-Z0-9]', '_', settings.OPENAI_MODEL).strip('_')
collection_name = f"knowledge_base_{model_slug}"[:63].rstrip('_')
collection_kwargs = {"name": collection_name, "metadata": {"hnsw:space": "cosine"}}
embedding_fn = get_embedding_function()
if embedding_fn:
    collection_kwargs["embedding_function"] = embedding_fn

try:
    collection = chroma_client.get_or_create_collection(**collection_kwargs)
except ValueError as e:
    if "Embedding function conflict" in str(e) or "already exists" in str(e):
        logger.warning(f"Embedding function conflict for collection '{collection_name}'. Re-creating collection for all-MiniLM-L6-v2 compatibility...")
        chroma_client.delete_collection(name=collection_name)
        collection = chroma_client.create_collection(**collection_kwargs)
    else:
        raise
logger.info(f"ChromaDB collection '{collection_name}' initialized at {CHROMA_PATH}")


def chunk_text(text: str, chunk_size: int = 500, chunk_overlap: int = 100) -> List[str]:
    """Split text into smaller chunks with overlap."""
    if not text:
        return []
        
    words = text.split()
    chunks = []
    
    # Simple chunking by words
    i = 0
    while i < len(words):
        chunk_words = words[i:i + chunk_size]
        chunks.append(" ".join(chunk_words))
        i += chunk_size - chunk_overlap
        if i >= len(words) or len(chunk_words) < chunk_size - chunk_overlap:
            break
            
    return chunks


def extract_pdf_text(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)

    text = ""

    for page in doc:
        text += page.get_text()

    doc.close()

    return text


def extract_and_chunk_pdf(pdf_path: str, chunk_size: int = 500, chunk_overlap: int = 100) -> List[str]:
    """Extract text from a PDF file and split it into smaller chunks."""
    text = extract_pdf_text(pdf_path)
    return chunk_text(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def extract_docx_text(docx_path: str) -> str:
    """Extract all text from a Word document (.docx/.doc)."""
    doc = docx.Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def extract_and_chunk_docx(docx_path: str, chunk_size: int = 500, chunk_overlap: int = 100) -> List[str]:
    """Extract text from a Word document and split it into smaller chunks."""
    text = extract_docx_text(docx_path)
    return chunk_text(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def extract_csv_text(csv_path: str) -> str:
    """Extract and format all text from a CSV file."""
    lines = []
    with open(csv_path, mode="r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        headers = next(reader, None)
        if not headers:
            return ""
        lines.append(" | ".join(headers))
        for row in reader:
            if any(cell.strip() for cell in row):
                row_str = " | ".join(
                    f"{headers[i]}: {cell.strip()}" if i < len(headers) and headers[i].strip() else cell.strip()
                    for i, cell in enumerate(row) if cell.strip()
                )
                lines.append(row_str)
    return "\n".join(lines)


def extract_and_chunk_csv(csv_path: str, chunk_size: int = 500, chunk_overlap: int = 100) -> List[str]:
    """Extract text from a CSV file and split it into smaller chunks."""
    text = extract_csv_text(csv_path)
    return chunk_text(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def get_embedding(text: str) -> List[float]:
    """Generate embedding for the text using SentenceTransformer (all-MiniLM-L6-v2) with LRU caching."""
    return get_cached_embedding(text)


def ingest_document(file_id: str, text: str) -> int:
    """Ingest a document: chunk it, embed it, and save to ChromaDB."""
    chunks = chunk_text(text)
    if not chunks:
        logger.warning(f"No content to ingest for file_id: {file_id}")
        return 0
        
    # Check if exact same document chunks are already in ChromaDB to avoid re-embedding
    try:
        existing_data = collection.get(where={"file_id": file_id}, include=["documents"])
        if existing_data and existing_data.get("documents") == chunks:
            logger.info(f"Document '{file_id}' chunks are unchanged in ChromaDB. Skipping embedding re-computation.")
            return len(chunks)
    except Exception as e:
        logger.warning(f"Could not check existing document cache for file_id {file_id}: {e}")

    # Delete existing chunks for this file if any before re-ingesting
    try:
        collection.delete(where={"file_id": file_id})
    except Exception as e:
        logger.warning(f"Error clearing existing chunks for file_id {file_id}: {e}")
        
    ids = [f"{file_id}_{idx}" for idx in range(len(chunks))]
    # get_embedding uses in-memory LRU cache + SentenceTransformer
    embeddings = [get_embedding(chunk) for chunk in chunks]
    metadatas = [{"file_id": file_id, "chunk_index": idx} for idx in range(len(chunks))]
    
    collection.add(
        ids=ids,
        embeddings=embeddings,
        documents=chunks,
        metadatas=metadatas
    )
    
    logger.info(f"Ingested {len(chunks)} chunks into ChromaDB collection '{collection.name}' for file_id: {file_id}")
    return len(chunks)


def ingest_pdf(file_id: str, pdf_path: str) -> int:
    """Ingest a PDF file: extract text using PyMuPDF, chunk it, embed it, and save to ChromaDB."""
    text = extract_pdf_text(pdf_path)
    return ingest_document(file_id, text)


def ingest_docx(file_id: str, docx_path: str) -> int:
    """Ingest a Word document (.docx/.doc): extract text, chunk it, embed it, and save to ChromaDB."""
    text = extract_docx_text(docx_path)
    return ingest_document(file_id, text)


def ingest_csv(file_id: str, csv_path: str) -> int:
    """Ingest a CSV file: extract text, chunk it, embed it, and save to ChromaDB."""
    text = extract_csv_text(csv_path)
    return ingest_document(file_id, text)


def ingest_file(file_id: str, file_path: str) -> int:
    """Unified file ingestion auto-router based on extension (.pdf, .docx, .doc, .csv, .txt, etc.)."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return ingest_pdf(file_id, file_path)
    elif ext in [".docx", ".doc"]:
        return ingest_docx(file_id, file_path)
    elif ext == ".csv":
        return ingest_csv(file_id, file_path)
    else:
        # Fallback for plain text, markdown, etc.
        with open(file_path, mode="r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return ingest_document(file_id, text)


def retrieve_context(query: str, top_k: int = 3) -> List[Dict[str, Any]]:
    """Retrieve top_k chunks relevant to the query from ChromaDB."""
    query_emb = get_embedding(query)
    
    try:
        results = collection.query(
            query_embeddings=[query_emb],
            n_results=top_k
        )
    except Exception as e:
        logger.error(f"Error querying ChromaDB: {e}")
        return []
        
    scored_chunks = []
    if not results or not results.get("documents") or len(results["documents"]) == 0:
        return []
        
    docs = results["documents"][0]
    metas = results["metadatas"][0]
    distances = results["distances"][0] if "distances" in results else [0.0] * len(docs)
    ids = results["ids"][0]
    
    for idx in range(len(docs)):
        # ChromaDB cosine space uses Cosine Distance (1 - cosine_similarity).
        # Convert back to cosine similarity for formatting: similarity = 1 - distance.
        distance = distances[idx]
        similarity = 1.0 - distance
        
        scored_chunks.append({
            "id": ids[idx],
            "file_id": metas[idx]["file_id"],
            "chunk_index": metas[idx]["chunk_index"],
            "content": docs[idx],
            "similarity": similarity
        })
        
    return scored_chunks


def generate_answer(query: str) -> Dict[str, Any]:
    """Retrieve relevant chunks and generate response using OpenAI."""
    contexts = retrieve_context(query, top_k=3)
    
    if not contexts:
        return {
            "answer": "No knowledge documents have been ingested yet. Please trigger ingestion via the webhook.",
            "context": []
        }
        
    context_str = "\n\n".join([f"[Source Chunk {idx+1}]:\n{ctx['content']}" for idx, ctx in enumerate(contexts)])
    
    prompt = f"""You are a helpful AI assistant. Use the following context from ingested documents to answer the user's question.
If the answer cannot be found or inferred from the context, state that the context does not contain the answer. Do not make up facts.

Context:
{context_str}

User Question:
{query}

Answer:"""

    client = get_openai_client()
    if not client:
        raise RuntimeError("OpenRouter/OpenAI API client is not configured. Please check your OPENROUTER_API_KEY / OPENAI_API_KEY settings.")

    try:
        response = client.chat.completions.create(
            model=settings.OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You are a helpful AI assistant answering questions based strictly on the provided knowledge context."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
        )
        answer = response.choices[0].message.content.strip()
        source = f"OpenAI/OpenRouter ({settings.OPENAI_MODEL})"
    except Exception as e:
        logger.error(f"Failed to generate answer from OpenAI/OpenRouter SDK: {e}")
        raise RuntimeError(f"LLM generation failed: {e}")
        
    return {
        "answer": answer,
        "source": source,
        "context": [
            {
                "file_id": ctx["file_id"],
                "chunk_index": ctx["chunk_index"],
                "content": ctx["content"],
                "similarity": round(ctx["similarity"], 4)
            }
            for ctx in contexts
        ]
    }
